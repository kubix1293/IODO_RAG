from support.ranking import effectiveness,total
from support.security import anonymize,anonymize_with_report,client_reference
from support.workflow import interpret_locally,validate_feedback
from support.worker import enrich_description,json_safe
import support.graph as support_graph
import support.learning as support_learning
import support.chat as support_chat
import support.knowledge_import as knowledge_import
from iodo_rag.chunking import detect_document_type,split_into_chunks
from iodo_rag.parsers import parse_docx
from docx import Document
import requests
import uuid

def test_ranking_weights(): assert total(1,1,1)==1
def test_effectiveness_has_prior(): assert effectiveness(1,0,0)<1
def test_anonymization():
    value=anonymize("Kontakt jan@example.com, tel. 501 502 503")
    assert "jan@example.com" not in value and "501 502 503" not in value
def test_external_privacy_redacts_pesel_and_credentials():
    value,categories=anonymize_with_report("Pacjent Jan Kowalski PESEL 44051401458, hasło: Tajne123, IP 10.2.3.4")
    assert "44051401458" not in value and "Jan Kowalski" not in value and "Tajne123" not in value and "10.2.3.4" not in value
    assert {"pesel_or_id","person","credential","ip"}.issubset(set(categories))
def test_client_reference_is_stable_pseudonym():
    assert client_reference(12,"secret")==client_reference(12,"secret")
    assert client_reference(12,"secret")!=client_reference(13,"secret")
def test_hybrid_llm_prefers_external(monkeypatch):
    monkeypatch.setattr(support_graph,"external_llm_answer",lambda prompt,runtime:"API")
    monkeypatch.setattr(support_graph,"local_llm_answer",lambda prompt,runtime:(_ for _ in ()).throw(AssertionError("fallback nie powinien ruszyć")))
    assert support_graph.hybrid_llm_answer("prompt",{"external_llm_enabled":1})==("API","external_api","")
def test_hybrid_llm_falls_back_to_ollama(monkeypatch):
    monkeypatch.setattr(support_graph,"external_llm_answer",lambda prompt,runtime:(_ for _ in ()).throw(TimeoutError("timeout")))
    monkeypatch.setattr(support_graph,"local_llm_answer",lambda prompt,runtime:"LOCAL")
    answer,provider,error=support_graph.hybrid_llm_answer("prompt",{"external_llm_enabled":1})
    assert answer=="LOCAL" and provider=="ollama_fallback" and "timeout" in error
def test_hybrid_llm_passes_approved_images_only_to_external(monkeypatch):
    seen={}
    monkeypatch.setattr(support_graph,"external_llm_answer",
                        lambda prompt,runtime,images:seen.setdefault("images",images) and "VISION")
    answer,provider,error=support_graph.hybrid_llm_answer(
        "prompt",{"external_llm_enabled":1},
        images=[{"id":"1","data_url":"data:image/png;base64,AA=="}])
    assert answer=="VISION" and provider=="external_api" and not error
    assert seen["images"][0]["id"]=="1"
def test_historical_case_image_sources_are_filtered_deduplicated_and_ordered():
    first=str(uuid.uuid4()); second=str(uuid.uuid4())
    sources=[
        {"kind":"documentation","id":"12"},
        {"kind":"historical_case","id":first},
        {"kind":"historical_case","id":"not-a-uuid"},
        {"kind":"historical_case","id":first},
        {"kind":"historical_case","id":second},
    ]
    assert support_graph.historical_case_ids_from_sources(sources)==[first,second]
def test_current_ticket_images_precede_historical_case_images(monkeypatch):
    current=[{"id":"ticket-image","origin":"current_ticket"}]
    monkeypatch.setattr(support_graph,"approved_ticket_images",lambda ticket_id:current)
    monkeypatch.setattr(support_graph,"historical_case_ids_from_sources",lambda sources:[])
    assert support_graph.approved_analysis_images({"ticket_id":"x","sources":[]})==current
def test_knowledge_curator_reuses_existing_solution(monkeypatch):
    monkeypatch.setattr(support_learning,"application_settings",lambda cur:{"external_llm_enabled":1})
    monkeypatch.setattr(support_learning,"hybrid_llm_answer",lambda prompt,runtime:(
        '{"action":"supplement","problem_id":7,"solution_id":9,"confidence":0.91,"reason":"ta sama przyczyna"}',
        "external_api",""))
    candidates=[{"problem_id":7,"problem_title":"Błąd usługi","problem_description":"Brak połączenia",
                 "solution_id":9,"solution_title":"Restart","solution_summary":"Uruchom usługę","solution_client_id":None}]
    decision,provider,error=support_learning.curate_knowledge(None,"Brak usługi","Restart usługi","Naprawa",candidates,"K-test")
    assert decision["action"]=="supplement" and decision["solution_id"]==9 and provider=="external_api" and not error
def test_knowledge_curator_rejects_unknown_ids(monkeypatch):
    monkeypatch.setattr(support_learning,"application_settings",lambda cur:{})
    monkeypatch.setattr(support_learning,"hybrid_llm_answer",lambda prompt,runtime:(
        '{"action":"duplicate","problem_id":999,"solution_id":888}',"external_api",""))
    try:
        support_learning.curate_knowledge(None,"Opis","Rozwiązanie","Tytuł",[],"K-test")
        assert False,"oczekiwano odrzucenia obcego ID"
    except ValueError:
        pass
def test_knowledge_curator_external_only_never_uses_local_fallback(monkeypatch):
    monkeypatch.setattr(support_learning,"application_settings",lambda cur:{"external_llm_enabled":1})
    monkeypatch.setattr(support_learning,"external_llm_answer",
                        lambda prompt,runtime:(_ for _ in ()).throw(TimeoutError("Qwen timeout")))
    monkeypatch.setattr(support_learning,"hybrid_llm_answer",
                        lambda *args,**kwargs:(_ for _ in ()).throw(AssertionError("fallback nie może ruszyć")))
    try:
        support_learning.curate_knowledge(
            None,"Opis","Rozwiązanie","Tytuł",[],"K-test",external_only=True)
        assert False,"oczekiwano błędu Qwen bez lokalnego fallbacku"
    except TimeoutError:
        pass
def test_curation_only_creates_case_for_new_knowledge():
    assert not support_learning.should_create_historical_case("duplicate")
    assert not support_learning.should_create_historical_case("supplement")
    assert support_learning.should_create_historical_case("new_solution")
    assert support_learning.should_create_historical_case("new_problem")
def test_effectiveness_counter_mapping():
    assert support_learning.effectiveness_counter("helped")=="success_count"
    assert support_learning.effectiveness_counter("partially_helped")=="partial_count"
    assert support_learning.effectiveness_counter("not_helped")=="failure_count"
def test_instruction_chunking_keeps_procedures_separate():
    text="# Restart usługi\n\n1. Otwórz panel usług.\n\n2. Uruchom usługę i sprawdź log.\n\n# Konfiguracja certyfikatu\n\n1. Otwórz magazyn certyfikatów.\n\n2. Zaimportuj certyfikat."
    assert detect_document_type(text)=="instruction"
    chunks=split_into_chunks(text,target_chars=100,overlap_chars=220)
    assert len(chunks)>=2
    assert all(not ("Restart usługi" in chunk["text"] and "Konfiguracja certyfikatu" in chunk["text"]) for chunk in chunks)
    assert any("Restart usługi" in chunk["text"] and "Uruchom usługę" in chunk["text"] for chunk in chunks)

def test_chunking_enforces_hard_limit_for_one_huge_legal_block():
    text="Art. 1. "+("bardzo długi tekst bez strukturalnego podziału. "*300)
    chunks=split_into_chunks(text,target_chars=700,overlap_chars=0)
    assert len(chunks)>1
    assert max(len(chunk["text"]) for chunk in chunks)<=700

def test_ai_chunk_groups_must_be_complete_contiguous_and_bounded():
    parts=[{"text":"A"*500,"title":"A"},{"text":"B"*500,"title":"B"},{"text":"C"*500,"title":"C"}]
    groups=knowledge_import._validate_groups([
        {"part_numbers":[1,2],"title":"Procedura","module":"ASW11","keywords":["RW"]},
        {"part_numbers":[3],"title":"Weryfikacja"},
    ],parts)
    assert groups[0]["part_numbers"]==[1,2]
    try:
        knowledge_import._validate_groups([
            {"part_numbers":[1,3],"title":"Błędna grupa"},
            {"part_numbers":[2],"title":"Powtórzenie"},
        ],parts)
        assert False,"oczekiwano odrzucenia nieciągłej propozycji"
    except ValueError:
        pass

def test_ai_chunk_proposals_preserve_source_text(monkeypatch):
    monkeypatch.setattr(knowledge_import,"split_into_chunks",lambda *args,**kwargs:[
        {"text":"Pierwszy krok.","title":"Krok 1","page_from":1,"page_to":1},
        {"text":"Drugi krok.","title":"Krok 2","page_from":1,"page_to":2},
    ])
    monkeypatch.setattr(knowledge_import,"_map_document",lambda *args,**kwargs:(
        {"document_title":"Instrukcja","modules":["ASW11"]},"external_api",""))
    monkeypatch.setattr(knowledge_import,"_group_batch",lambda parts,*args,**kwargs:([
        {"part_numbers":[1,2],"title":"Wydanie RW","module":"ASW11","operation":"wydanie",
         "content_type":"procedure","keywords":["RW","bufor"]}
    ],"external_api",""))
    proposals,document_map,provider,error=knowledge_import.propose_chunks(
        "tekst",[],"ASW",{"llm_response_tokens":1200})
    assert proposals[0]["chunk_text"]=="Pierwszy krok.\n\nDrugi krok."
    assert proposals[0]["metadata"]["page_from"]==1
    assert document_map["modules"]==["ASW11"] and provider=="external_api" and not error
def test_docx_parser_preserves_heading_style(tmp_path):
    path=tmp_path/"instruction.docx"; document=Document()
    document.add_heading("Procedura aktualizacji",level=1)
    document.add_paragraph("1. Wykonaj kopię zapasową.")
    document.save(path)
    text,_=parse_docx(path)
    assert "# Procedura aktualizacji" in text
def test_support_prompt_is_technical_not_legal():
    prompt=support_graph.build_technical_support_prompt({
        "client_ref":"K-test","effective_description":"ERR-1234 podczas zapisu w module ASW",
        "sources":[{"kind":"documentation","id":"1","chunk_text":"Sprawdź usługę ZSIMED i jej log."}],
    })
    assert "kolejnych numerowanych kroków" in prompt and "co zrobić, dlaczego" in prompt
    assert "Nie twórz stylu prawnego" in prompt and "Nie wypisuj numerów materiałów" in prompt
    assert "Nie używaj składni Markdown" in prompt
def test_task_prompt_is_short_and_redirects_to_consultation():
    prompt=support_graph.build_technical_support_prompt({
        "client_ref":"K-test","effective_description":"Aktualizacja modułu ZZL",
        "recognized":{"issue_kind":"task"},"sources":[],
    })
    assert "W ramach tego zadania pamiętaj o:" in prompt
    assert "Nie diagnozuj przyczyny" in prompt
    assert "Konsultacja AI" in prompt
def test_llm_markdown_emphasis_is_removed():
    assert support_graph.plain_text_response("## Diagnoza\n**Błąd**\n* krok")=="Diagnoza\nBłąd\n- krok"
def test_support_prompt_uses_titles_full_chunks_and_global_budget():
    full_chunk="A"*1600
    prompt=support_graph.build_technical_support_prompt({
        "client_ref":"K-test","effective_description":"ERR-1234",
        "sources":[
            {"kind":"documentation","title":"Restart usługi","context_role":"match","chunk_text":full_chunk},
            {"kind":"documentation","title":"Restart usługi","context_role":"neighbor","chunk_text":"kolejny krok"},
            *[{"kind":"documentation","title":"Długi dokument","chunk_text":"B"*5000} for _ in range(8)],
        ],
    })
    materials=prompt.split("MATERIAŁY TECHNICZNE:\n",1)[1]
    assert full_chunk in materials and "TYTUŁ: Restart usługi" in materials
    assert "fragment sąsiedni procedury" in materials
    assert len(materials)<=support_graph.LLM_CONTEXT_MAX_CHARS+1
def test_documentation_match_is_expanded_with_adjacent_chunks():
    class Cursor:
        def execute(self,query,params):
            assert params==(7,2,4,2,11)
        def fetchall(self):
            return [
                {"id":"30","kind":"documentation","document_id":7,"chunk_index":2,"chunk_text":"przed","title":"Instrukcja"},
                {"id":"32","kind":"documentation","document_id":7,"chunk_index":4,"chunk_text":"po","title":"Instrukcja"},
            ]
    selected=[{"id":"31","kind":"documentation","document_id":7,"chunk_index":3,
               "chunk_text":"trafienie","title":"Instrukcja","rerank_score":0.9}]
    expanded=support_graph.expand_documentation_neighbors(
        selected,{"program_id":2,"client_id":11},Cursor())
    assert [row["chunk_index"] for row in expanded]==[2,3,4]
    assert [row["context_role"] for row in expanded]==["neighbor","match","neighbor"]
def test_feedback_validation():
    assert validate_feedback("not_helped","")[0]=="incomplete"
    assert validate_feedback("helped","nadal nie pomogło")[0]=="suspicious"
    assert validate_feedback("helped","wykonano poprawnie")[0]=="consistent"
def test_interpretation():
    result=interpret_locally("Błąd ERR-1234 w wersji 2.4.1 podczas zapisu")
    assert result["error_code"] and result["version"]=="2.4.1" and result["issue_kind"]=="problem"
def test_task_interpretation_does_not_require_error_or_version():
    result=interpret_locally("Aktualizacja modułu ZZL i przygotowanie pakietu FAKTURY")
    assert result["issue_kind"]=="task" and result["missing"]==[]
def test_non_finite_scores_are_json_safe(): assert json_safe({"score":float("nan")})=={"score":None}
def test_clarification_answers_reach_model_context():
    enriched=enrich_description("Problem z usługą",{"error_code":"brak","version":"4.2"})
    assert "error_code: brak" in enriched and "version: 4.2" in enriched

def test_state_graph_runs_both_db_agents(monkeypatch):
    monkeypatch.setattr(support_graph,"history_agent_node",lambda state:{"history_candidates":[{"kind":"historical_case","chunk_text":"historia"}]})
    monkeypatch.setattr(support_graph,"documentation_agent_node",lambda state:{"documentation_candidates":[{"kind":"documentation","chunk_text":"instrukcja"}]})
    monkeypatch.setattr(support_graph,"reranking_node",lambda state:{"sources":state["history_candidates"]+state["documentation_candidates"],"step":"answer_generation"})
    monkeypatch.setattr(support_graph,"answer_node",lambda state:{"proposed_answer":"gotowe","status":"awaiting_problem_decision","step":"problem_decision"})
    graph=support_graph.build_graph(None)
    result=graph.invoke({"ticket_id":"test","client_id":1,"program_id":1,"description":"ERR-1234 wersja 1.2.3","answers":{},"history_candidates":[],"documentation_candidates":[]})
    assert {row["kind"] for row in result["sources"]}=={"historical_case","documentation"}
    assert result["proposed_answer"]=="gotowe"

def test_consultation_retrieves_again_for_each_correction(monkeypatch):
    seen={}
    def history(state):
        seen["query"]=state["effective_description"]
        return {"history_candidates":[]}
    monkeypatch.setattr(support_chat,"history_agent_node",history)
    monkeypatch.setattr(support_chat,"documentation_agent_node",lambda state:{"documentation_candidates":[]})
    monkeypatch.setattr(support_chat,"reranking_node",lambda state:{"sources":[{"kind":"documentation","title":"Instrukcja","chunk_text":"Wykonaj krok A"}]})
    monkeypatch.setattr(support_chat,"application_settings",lambda:{})
    monkeypatch.setattr(support_chat,"hybrid_llm_answer",lambda prompt,runtime:(seen.setdefault("prompt",prompt) and "Odpowiedź","external_api",""))
    answer,provider,error,sources=support_chat.answer_consultation(
        question="Sprostowanie: błąd występuje przy eksporcie",conversation=[{"role":"user","content":"Problem przy zapisie"}],
        ticket_description="ERR-1234 w ASW",program_id=2,client_id=7,client_ref="K-test")
    assert "Sprostowanie: błąd występuje przy eksporcie" in seen["query"]
    assert "HISTORIA ROZMOWY" in seen["prompt"] and "Wykonaj krok A" in seen["prompt"]
    assert answer=="Odpowiedź" and provider=="external_api" and not error and len(sources)==1

def test_reranker_timeout_uses_lexical_fallback(monkeypatch):
    monkeypatch.setattr(support_graph.requests,"post",lambda *args,**kwargs:(_ for _ in ()).throw(requests.Timeout("timeout")))
    scores=support_graph.rerank("ERR-1234 eksport ASW",["Instrukcja ERR-1234 dla eksportu","Ogólne logowanie"])
    assert scores[0]>scores[1]
